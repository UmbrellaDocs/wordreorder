#!/usr/bin/env python3
import argparse
import sys
import copy
from pathlib import Path
import yaml
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from tqdm import tqdm
import colorama
from colorama import Fore, Style, init as colorama_init

def is_heading_style(paragraph, max_level):
    if not paragraph or not paragraph.style or not paragraph.style.name:
        return None, None
    style_name = paragraph.style.name
    style_name_lower = style_name.strip().lower()
    if style_name_lower.startswith("heading "):
        try:
            level = int(style_name.split(" ")[-1])
            if 1 <= level <= max_level:
                return level, paragraph.text
        except (ValueError, IndexError):
            return None, None
    return None, None

def parse_document_structure(doc_path, max_level, include_elements=False):
    try:
        document = Document(doc_path)
    except PackageNotFoundError:
        print(f"{Fore.RED}Error: Could not find or open Word document: {doc_path}{Style.RESET_ALL}", file=sys.stderr)
        return None, None
    except Exception as e:
        print(f"{Fore.RED}Error: Failed to parse Word document '{doc_path}': {e}{Style.RESET_ALL}", file=sys.stderr)
        return None, None

    sections = []
    current_section_elements = [] if include_elements else None
    current_heading_info = {'text': '__PREAMBLE__', 'level': 0}

    print(f"{Fore.BLUE}Parsing source document: {doc_path}...{Style.RESET_ALL}")
    iterator = tqdm(document.element.body, desc="Scanning Elements", unit="element")
    for element in iterator:
        is_paragraph = element.tag.endswith('}p')
        para_obj = None
        if is_paragraph:
            for p in document.paragraphs:
                if p._element is element:
                    para_obj = p
                    break
            if para_obj is not None:
                level, heading_text = is_heading_style(para_obj, max_level)
                if level is not None:
                    if current_heading_info['level'] > 0 or (current_section_elements is not None and current_section_elements):
                         sections.append({
                             'text': current_heading_info['text'].strip(),
                             'level': current_heading_info['level'],
                             'elements': current_section_elements if include_elements else None
                         })
                    current_heading_info = {'text': heading_text, 'level': level}
                    current_section_elements = [element] if include_elements else None
                    continue
        if include_elements and current_section_elements is not None:
            current_section_elements.append(element)

    if current_heading_info['level'] > 0 or (current_section_elements is not None and current_section_elements):
        sections.append({
            'text': current_heading_info['text'].strip(),
            'level': current_heading_info['level'],
            'elements': current_section_elements if include_elements else None
        })

    if include_elements and sections and sections[0]['text'] == '__PREAMBLE__' and not sections[0]['elements']:
        sections.pop(0)

    num_headings = len([s for s in sections if s['level'] > 0])
    print(f"{Fore.GREEN}Found {num_headings} heading sections (plus potential preamble).{Style.RESET_ALL}")
    return sections, document


def build_nested_toc(flat_sections):
    nested_toc = []
    stack = []
    for section in flat_sections:
        level = section['level']
        if level == 0:
             continue
        node = {'heading': section['text'], 'level': level}
        while stack and stack[-1][0] >= level:
            stack.pop()
        if not stack:
            nested_toc.append(node)
        else:
            parent_node_list = stack[-1][1]
            if 'children' not in parent_node_list[-1]:
                parent_node_list[-1]['children'] = []
            parent_node_list[-1]['children'].append(node)
        if 'children' not in node:
             node['children'] = []
        stack.append((level, [node] if not stack else stack[-1][1][-1]['children']))

    def cleanup_empty_children(node_list):
        for node in node_list:
            if 'children' in node:
                if node['children']:
                    cleanup_empty_children(node['children'])
                if not node['children']:
                    del node['children']
    cleanup_empty_children(nested_toc)
    return nested_toc

def load_toc_config_for_reorganize(toc_path):
    try:
        with open(toc_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
        if not isinstance(config, dict) or 'toc' not in config:
            print(f"{Fore.RED}Error: YAML file '{toc_path}' must contain a top-level 'toc' key.{Style.RESET_ALL}", file=sys.stderr)
            return None
        toc_data = config['toc']
        if not isinstance(toc_data, list):
             print(f"{Fore.RED}Error: The 'toc' key in '{toc_path}' must contain a list.{Style.RESET_ALL}", file=sys.stderr)
             return None

        flat_headings = []
        def extract_headings(node_list):
            for item in node_list:
                if isinstance(item, str):
                    flat_headings.append(item)
                elif isinstance(item, dict) and 'heading' in item:
                    flat_headings.append(item['heading'])
                    if 'children' in item and isinstance(item['children'], list):
                        extract_headings(item['children'])
                else:
                    print(f"{Fore.YELLOW}Warning: Skipping unrecognized item in TOC: {item}{Style.RESET_ALL}", file=sys.stderr)

        extract_headings(toc_data)

        if not flat_headings:
             print(f"{Fore.RED}Error: No valid heading entries found under 'toc' key in '{toc_path}'.{Style.RESET_ALL}", file=sys.stderr)
             return None

        return flat_headings

    except FileNotFoundError:
        print(f"{Fore.RED}Error: TOC file not found: {toc_path}{Style.RESET_ALL}", file=sys.stderr)
        return None
    except yaml.YAMLError as e:
        print(f"{Fore.RED}Error: Could not parse YAML file '{toc_path}': {e}{Style.RESET_ALL}", file=sys.stderr)
        return None
    except Exception as e:
        print(f"{Fore.RED}Error: An unexpected error occurred loading TOC '{toc_path}': {e}{Style.RESET_ALL}", file=sys.stderr)
        return None


def run_generate(args):
    print(f"{Fore.CYAN}--- Generating TOC Structure ---{Style.RESET_ALL}")
    flat_sections, _ = parse_document_structure(args.input, args.max_level, include_elements=False)
    if flat_sections is None:
        return False

    nested_toc_structure = build_nested_toc(flat_sections)
    yaml_data = {'toc': nested_toc_structure}

    try:
        print(f"{Fore.BLUE}Writing generated TOC structure to: {args.output}...{Style.RESET_ALL}")
        with open(args.output, 'w', encoding='utf-8') as f:
            yaml.safe_dump(yaml_data, f, sort_keys=False, allow_unicode=True, indent=2)
        print(f"{Fore.GREEN}TOC generation complete.{Style.RESET_ALL}")
        print(f"You can now manually edit '{args.output}' and use it with the 'reorganize' command.")
        return True
    except Exception as e:
        print(f"{Fore.RED}Error: Failed to write YAML file '{args.output}': {e}{Style.RESET_ALL}", file=sys.stderr)
        return False

def run_reorganize(args):
    print(f"{Fore.CYAN}--- Reorganizing Document ---{Style.RESET_ALL}")

    print(f"{Fore.BLUE}Loading target TOC structure from: {args.toc}...{Style.RESET_ALL}")
    target_toc_headings = load_toc_config_for_reorganize(args.toc)
    if target_toc_headings is None:
        return False

    source_sections_data, source_doc_obj = parse_document_structure(args.input, args.max_level, include_elements=True)
    if source_sections_data is None:
        return False

    source_sections_map = {}
    duplicate_headings = set()
    headings_seen = set()
    preamble_section = None
    for section in source_sections_data:
        heading_text = section['text']
        if heading_text == '__PREAMBLE__':
            preamble_section = section
            continue
        norm_heading = heading_text
        if norm_heading in headings_seen and norm_heading not in duplicate_headings:
             print(f"{Fore.YELLOW}Warning: Duplicate heading text found in source: '{heading_text}'. Using first instance.{Style.RESET_ALL}", file=sys.stderr)
             duplicate_headings.add(norm_heading)
             continue
        elif norm_heading not in headings_seen:
            source_sections_map[norm_heading] = section
            headings_seen.add(norm_heading)

    print(f"{Fore.BLUE}Comparing source sections with target TOC...{Style.RESET_ALL}")
    sections_to_write = []
    found_source_headings = set(source_sections_map.keys())
    target_toc_set = set(target_toc_headings)

    missing_in_source = []
    for heading in target_toc_headings:
         norm_heading = heading
         if norm_heading not in found_source_headings:
              missing_in_source.append(heading)

    unmatched_in_source_headings = found_source_headings - target_toc_set

    if missing_in_source:
        msg = f"Headings in TOC but not found in source (or duplicates skipped): {', '.join(missing_in_source)}"
        if args.missing == 'error':
            print(f"{Fore.RED}Error: {msg} (Policy: error){Style.RESET_ALL}", file=sys.stderr)
            return False
        elif args.missing == 'warn':
            print(f"{Fore.YELLOW}Warning: {msg} (Policy: warn, skipping){Style.RESET_ALL}", file=sys.stderr)

    unmatched_sections_to_append = []
    if unmatched_in_source_headings:
        unmatched_original_text = [source_sections_map[h]['text'] for h in unmatched_in_source_headings]
        msg = f"Headings in source (first instance) but not in TOC: {', '.join(unmatched_original_text)}"
        if args.unmatched == 'delete':
            print(f"{Fore.YELLOW}Info: {msg} (Policy: delete, discarding){Style.RESET_ALL}", file=sys.stderr)
        else:
             policy_desc = "append, adding to end"
             color = Fore.BLUE
             if args.unmatched == 'warn':
                  policy_desc = "warn, will append to end"
                  color = Fore.YELLOW
             print(f"{color}Info: {msg} (Policy: {policy_desc}){Style.RESET_ALL}", file=sys.stderr)
             for heading in unmatched_in_source_headings:
                 unmatched_sections_to_append.append(source_sections_map[heading])

    print(f"{Fore.BLUE}Planning writing order...{Style.RESET_ALL}")
    if preamble_section and preamble_section.get('elements'):
        print("Including preamble content at the beginning.")
        sections_to_write.append(preamble_section)
    processed_target_headings = set()
    for heading in target_toc_headings:
        norm_heading = heading
        if norm_heading in source_sections_map and norm_heading not in processed_target_headings:
            sections_to_write.append(source_sections_map[norm_heading])
            processed_target_headings.add(norm_heading)
            if args.verbose:
                print(f"  - Scheduling section: '{source_sections_map[norm_heading]['text']}'")
        elif norm_heading not in source_sections_map and args.verbose:
            print(f"  - Skipping missing/duplicate section from TOC: '{heading}'")

    if unmatched_sections_to_append:
        print(f"Scheduling {len(unmatched_sections_to_append)} unmatched sections for appending.")
        sections_to_write.extend(unmatched_sections_to_append)


    if not sections_to_write:
        print(f"{Fore.YELLOW}Warning: No sections identified to write to the output document.{Style.RESET_ALL}", file=sys.stderr)

    print(f"{Fore.BLUE}Creating new document: {args.output}...{Style.RESET_ALL}")
    new_document = Document()

    total_elements_to_copy = sum(len(section.get('elements', [])) for section in sections_to_write)
    if total_elements_to_copy == 0 and args.verbose:
        print("Note: No actual elements found to copy across scheduled sections.")

    copy_errors = 0
    with tqdm(total=total_elements_to_copy, desc="Writing Sections", unit="element", disable=not total_elements_to_copy) as pbar:
        for section in sections_to_write:
            if not section.get('elements'):
                continue
            if args.verbose:
                 print(f"  - Writing section: '{section['text']}' ({len(section['elements'])} elements)")

            for element in section['elements']:
                try:
                    element_copy = copy.deepcopy(element)
                    new_document.element.body.append(element_copy)
                    pbar.update(1)
                except Exception as e:
                    copy_errors += 1
                    if copy_errors <= 5:
                        print(f"\n{Fore.YELLOW}Warning: Failed to copy an element in section '{section['text']}'. Error: {e}{Style.RESET_ALL}", file=sys.stderr)
                    if copy_errors == 6:
                         print(f"\n{Fore.YELLOW}Warning: Further element copy errors will be suppressed...{Style.RESET_ALL}", file=sys.stderr)
                    pbar.update(1)

    if copy_errors > 0:
        print(f"\n{Fore.YELLOW}Warning: Encountered {copy_errors} errors during element copying. Review output carefully.{Style.RESET_ALL}", file=sys.stderr)

    try:
        print(f"{Fore.BLUE}Saving reorganized document to: {args.output}...{Style.RESET_ALL}")
        new_document.save(args.output)
        print(f"{Fore.GREEN}Reorganization complete!{Style.RESET_ALL}")
        return True
    except Exception as e:
        print(f"{Fore.RED}Error: Failed to save the reorganized document '{args.output}': {e}{Style.RESET_ALL}", file=sys.stderr)
        return False

def main():
    colorama_init(autoreset=True)

    parser = argparse.ArgumentParser(
        description="Reorganize a Word document based on a YAML Table of Contents, or generate a TOC YAML from a document.",
        formatter_class=argparse.RawTextHelpFormatter
    )
    subparsers = parser.add_subparsers(dest="command", required=True, help="Available commands")

    parser_gen = subparsers.add_parser(
        "generate", help="Generate a structured TOC YAML file from a Word document.", description="Reads a Word document and creates a YAML file representing its heading structure (up to --max-level). This YAML can then be edited and used with the 'reorganize' command."
    )
    parser_gen.add_argument( "-i", "--input", required=True, type=Path, help="Path to the source Word document (.docx)." )
    parser_gen.add_argument( "-o", "--output", required=True, type=Path, help="Path to save the generated YAML file (.yaml)." )
    parser_gen.add_argument( "-l", "--max-level", type=int, default=6, help="Maximum heading level to include in the generated TOC. Default: 6." )
    parser_gen.add_argument( "-v", "--verbose", action="store_true", help="Increase output verbosity." )

    parser_reorg = subparsers.add_parser(
        "reorganize", help="Reorganize a Word document based on a TOC YAML file.", description="Restructures a Word document according to the heading order specified in a YAML file."
    )
    parser_reorg.add_argument( "-i", "--input", required=True, type=Path, help="Path to the source Word document (.docx)." )
    parser_reorg.add_argument( "-t", "--toc", required=True, type=Path, help="Path to the YAML file defining the target TOC structure.\n..." )
    parser_reorg.add_argument( "-o", "--output", required=True, type=Path, help="Path to save the reorganized Word document (.docx)." )
    parser_reorg.add_argument( "-l", "--max-level", type=int, default=6, help="Maximum heading level used to identify sections in the source document. Default: 6." )
    parser_reorg.add_argument( "-u", "--unmatched", choices=['append', 'delete', 'warn'], default='append', help="How to handle source sections not in TOC: append (default), delete, warn (and append)." )
    parser_reorg.add_argument( "-m", "--missing", choices=['error', 'warn', 'ignore'], default='warn', help="How to handle TOC sections not in source: error, warn (default), ignore." )
    parser_reorg.add_argument( "-v", "--verbose", action="store_true", help="Increase output verbosity." )

    args = parser.parse_args()

    if not args.input.is_file():
        print(f"{Fore.RED}Error: Input Word file not found: {args.input}{Style.RESET_ALL}", file=sys.stderr)
        sys.exit(1)

    if args.command == "reorganize":
        if not args.toc.is_file():
            print(f"{Fore.RED}Error: TOC YAML file not found: {args.toc}{Style.RESET_ALL}", file=sys.stderr)
            sys.exit(1)
        if args.output.is_dir():
             print(f"{Fore.RED}Error: Reorganize output path is a directory: {args.output}{Style.RESET_ALL}", file=sys.stderr)
             sys.exit(1)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        success = run_reorganize(args)

    elif args.command == "generate":
        if args.output.is_dir():
             print(f"{Fore.RED}Error: Generate output path is a directory: {args.output}{Style.RESET_ALL}", file=sys.stderr)
             sys.exit(1)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        success = run_generate(args)

    else:
        print(f"{Fore.RED}Error: Unknown command '{args.command}'{Style.RESET_ALL}", file=sys.stderr)
        parser.print_help()
        sys.exit(1)

    if success:
        sys.exit(0)
    else:
        print(f"\n{Fore.RED}Process failed.{Style.RESET_ALL}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
